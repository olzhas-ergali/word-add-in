"""
API для работы с шаблонами
Загрузка шаблонов из Printable Forms API и автоматическое сохранение параметров в БД
"""

from fastapi import APIRouter, HTTPException
from typing import List
from uuid import UUID
import re
from io import BytesIO
from docx import Document

from app.models.document import PfDocument
from app.services.pf_api_service import PfApiService
from app.services.database_service import db_service

router = APIRouter()


async def extract_variables_from_docx(docx_bytes: bytes) -> List[dict]:
    """
    Извлечь переменные из Word документа
    
    Поддерживаемые форматы:
    - {VARIABLE_NAME}
    - {{VARIABLE_NAME}}
    - [VARIABLE_NAME]
    - ${VARIABLE_NAME}
    """
    doc = Document(BytesIO(docx_bytes))
    variables = set()
    
    # Паттерны для поиска переменных
    patterns = [
        r'\{([A-Za-z_][A-Za-z0-9_]*)\}',           # {VAR}
        r'\{\{([A-Za-z_][A-Za-z0-9_]*)\}\}',       # {{VAR}}
        r'\[([A-Za-z_][A-Za-z0-9_]*)\]',           # [VAR]
        r'\$\{([A-Za-z_][A-Za-z0-9_]*)\}',         # ${VAR}
    ]
    
    # Извлекаем из параграфов
    for para in doc.paragraphs:
        text = para.text
        for pattern in patterns:
            matches = re.findall(pattern, text)
            variables.update(matches)
    
    # Извлекаем из таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                for pattern in patterns:
                    matches = re.findall(pattern, text)
                    variables.update(matches)
    
    # Извлекаем из Custom Properties
    try:
        for prop in doc.core_properties.custom_properties:
            variables.add(prop.name)
    except:
        pass
    
    return [{"name": var, "value": ""} for var in sorted(variables)]


async def save_variables_to_db(document_id: str, document_name: str, variables: List[dict]):
    """
    Сохранить переменные в БД
    """
    if not db_service.pool:
        await db_service.connect()
    
    async with db_service.pool.acquire() as conn:
        for var in variables:
            # Вставляем параметр, если уже существует - пропускаем
            await conn.execute("""
                INSERT INTO contract_parameters 
                (contract_id, param_name, param_value, description)
                VALUES ($1, $2, $3, $4)
                ON CONFLICT (contract_id, param_name) DO NOTHING
            """, 
            document_id,
            var['name'],
            var.get('value', ''),
            f"Автоматически извлечено из {document_name}"
            )


@router.get("/list", response_model=List[PfDocument])
async def get_template_list():
    """
    Получить список шаблонов из Printable Forms API
    
    Returns:
        List[PfDocument] - список доступных шаблонов
    """
    pf_service = PfApiService()
    
    try:
        documents = await pf_service.get_template_files(token=None)
        return documents
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Ошибка получения списка шаблонов: {str(e)}"
        )


@router.post("/download/{document_id}")
async def download_and_save_template(document_id: UUID):
    """
    Загрузить шаблон из API и автоматически сохранить параметры в БД
    
    Args:
        document_id: UUID документа
        
    Returns:
        Информация о загруженном документе и сохраненных параметрах
    """
    pf_service = PfApiService()
    
    try:
        # 1. Загрузить документ из API
        docx_bytes = await pf_service.get_pf_document(document_id, token=None)
        
        if not docx_bytes:
            raise HTTPException(status_code=404, detail="Документ не найден в API")
        
        # 2. Получить метаданные документа
        documents = await pf_service.get_template_files(token=None)
        document = next((d for d in documents if d.document_id == document_id), None)
        
        if not document:
            raise HTTPException(status_code=404, detail="Метаданные документа не найдены")
        
        document_name = document.file_name
        contract_id = f"TEMPLATE-{str(document_id)[:8]}"
        
        # 3. Извлечь переменные из документа
        variables = await extract_variables_from_docx(docx_bytes)
        
        # 4. Сохранить переменные в БД
        if variables:
            await save_variables_to_db(contract_id, document_name, variables)
        
        # 5. Вернуть документ и информацию о сохраненных параметрах
        return {
            "success": True,
            "document": {
                "id": str(document_id),
                "name": document_name,
                "contract_id": contract_id,
                "file_size": len(docx_bytes)
            },
            "variables": {
                "found": len(variables),
                "saved_to_db": len(variables),
                "contract_id": contract_id,
                "list": variables
            },
            "message": f"Документ загружен. Найдено и сохранено {len(variables)} параметров в БД"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Ошибка обработки документа: {str(e)}"
        )


@router.get("/download/{document_id}/file")
async def download_template_file(document_id: UUID):
    """
    Скачать только файл шаблона (без сохранения параметров)
    
    Args:
        document_id: UUID документа
        
    Returns:
        Файл документа
    """
    from fastapi.responses import StreamingResponse
    
    pf_service = PfApiService()
    
    try:
        docx_bytes = await pf_service.get_pf_document(document_id, token=None)
        
        if not docx_bytes:
            raise HTTPException(status_code=404, detail="Документ не найден")
        
        return StreamingResponse(
            BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=template_{document_id}.docx"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Ошибка загрузки файла: {str(e)}"
        )


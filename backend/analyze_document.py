#!/usr/bin/env python3
"""
Расширенный анализ Word документа для поиска всех возможных переменных
"""

import re
from docx import Document
from docx.oxml import parse_xml
from pathlib import Path
from collections import Counter
import sys


def analyze_document(docx_path: str):
    """Полный анализ документа"""
    
    doc = Document(docx_path)
    
    print("=" * 70)
    print("📊 АНАЛИЗ ДОКУМЕНТА")
    print("=" * 70)
    print(f"Файл: {docx_path}")
    print(f"Параграфов: {len(doc.paragraphs)}")
    print(f"Таблиц: {len(doc.tables)}")
    print()
    
    # 1. Поиск всех возможных паттернов
    print("🔍 ПОИСК ПЕРЕМЕННЫХ В РАЗНЫХ ФОРМАТАХ")
    print("-" * 70)
    
    all_text = []
    
    # Собираем весь текст
    for para in doc.paragraphs:
        all_text.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text.append(cell.text)
    
    full_text = "\n".join(all_text)
    
    # Различные паттерны
    patterns = {
        "Фигурные скобки {VAR}": r'\{([^}]+)\}',
        "Двойные фигурные {{VAR}}": r'\{\{([^}]+)\}\}',
        "Квадратные скобки [VAR]": r'\[([^\]]+)\]',
        "Двойные квадратные [[VAR]]": r'\[\[([^\]]+)\]\]',
        "Dollar ${VAR}": r'\$\{([^}]+)\}',
        "Процент %VAR%": r'%([^%]+)%',
        "Знак подчеркивания _VAR_": r'_([^_]+)_',
        "Точки ...VAR...": r'\.\.\.([^.]+)\.\.\.',
        "Угловые <VAR>": r'<([^>]+)>',
        "Числа в скобках (1), (2)": r'\((\d+)\)',
    }
    
    found_any = False
    all_variables = {}
    
    for pattern_name, pattern in patterns.items():
        matches = re.findall(pattern, full_text)
        if matches:
            # Отфильтруем очень длинные (вероятно не переменные)
            filtered = [m for m in matches if len(m) < 100]
            if filtered:
                found_any = True
                all_variables[pattern_name] = filtered
                print(f"\n✅ {pattern_name}: найдено {len(filtered)}")
                
                # Показываем первые 10
                for i, match in enumerate(filtered[:10], 1):
                    preview = match[:50] + "..." if len(match) > 50 else match
                    print(f"   {i}. {preview}")
                
                if len(filtered) > 10:
                    print(f"   ... и еще {len(filtered) - 10}")
    
    if not found_any:
        print("❌ Переменные не найдены ни в одном формате")
    
    print()
    print("=" * 70)
    
    # 2. Показываем начало документа
    print("📄 НАЧАЛО ДОКУМЕНТА (первые 20 строк):")
    print("-" * 70)
    for i, para in enumerate(doc.paragraphs[:20], 1):
        if para.text.strip():
            text = para.text[:100]
            print(f"{i:2d}. {text}")
    
    print()
    print("=" * 70)
    
    # 3. Анализ таблиц
    if doc.tables:
        print("📊 ПЕРВАЯ ТАБЛИЦА:")
        print("-" * 70)
        first_table = doc.tables[0]
        for i, row in enumerate(first_table.rows[:5], 1):
            row_text = " | ".join([cell.text[:30] for cell in row.cells[:4]])
            print(f"Строка {i}: {row_text}")
        
        if len(first_table.rows) > 5:
            print(f"... и еще {len(first_table.rows) - 5} строк")
        
        print()
        print("=" * 70)
    
    # 4. Проверка на Content Controls
    print("🎛️  ПОИСК CONTENT CONTROLS:")
    print("-" * 70)
    
    content_controls_found = False
    for para in doc.paragraphs[:50]:  # Проверяем первые 50 параграфов
        for run in para.runs:
            # Проверяем XML на наличие Content Controls
            if 'w:sdt' in str(run.element.xml):
                content_controls_found = True
                break
        if content_controls_found:
            break
    
    if content_controls_found:
        print("✅ Документ содержит Content Controls")
        print("   (Это специальные элементы Word для ввода данных)")
    else:
        print("❌ Content Controls не найдены")
    
    print()
    print("=" * 70)
    
    # 5. Поиск повторяющихся слов (возможно это метки)
    print("🔤 ЧАСТО ВСТРЕЧАЮЩИЕСЯ СЛОВА (возможные метки):")
    print("-" * 70)
    
    # Разбиваем на слова
    words = re.findall(r'\b[А-ЯЁ][а-яё]+\b', full_text)
    word_counts = Counter(words)
    
    # Показываем топ-20 слов длиннее 5 символов
    common_words = [(word, count) for word, count in word_counts.most_common(50) 
                    if len(word) > 5 and count > 5]
    
    if common_words:
        for i, (word, count) in enumerate(common_words[:20], 1):
            print(f"{i:2d}. {word:30s} - {count} раз")
    
    print()
    print("=" * 70)
    
    # 6. Статистика
    print("📈 ОБЩАЯ СТАТИСТИКА:")
    print("-" * 70)
    print(f"Всего символов: {len(full_text):,}")
    print(f"Всего слов: {len(full_text.split()):,}")
    print(f"Всего строк: {len(all_text):,}")
    print()
    
    # 7. Сохраняем образец текста
    sample_path = Path("variables_export/document_sample.txt")
    sample_path.parent.mkdir(exist_ok=True)
    
    with open(sample_path, 'w', encoding='utf-8') as f:
        f.write("ОБРАЗЕЦ ТЕКСТА ДОКУМЕНТА\n")
        f.write("=" * 70 + "\n\n")
        for i, para in enumerate(doc.paragraphs[:100], 1):
            if para.text.strip():
                f.write(f"{i}. {para.text}\n\n")
    
    print(f"💾 Образец текста сохранен: {sample_path}")
    print()
    
    # 8. Если нашли переменные, предлагаем создать БД
    if all_variables:
        print("=" * 70)
        print("✅ НАЙДЕНЫ ПЕРЕМЕННЫЕ!")
        print("=" * 70)
        print()
        print("Хотите создать базу данных из этих переменных?")
        print("Для этого используйте переменные из категории,")
        print("которая больше всего подходит вашему документу.")
        print()
        
        # Сохраняем все найденные переменные
        import json
        all_vars_path = Path("variables_export/all_found_variables.json")
        with open(all_vars_path, 'w', encoding='utf-8') as f:
            json.dump(all_variables, f, ensure_ascii=False, indent=2)
        print(f"💾 Все найденные переменные: {all_vars_path}")
    

def main():
    docx_path = "../ДДУ Шымкент.docx"
    
    if not Path(docx_path).exists():
        print(f"❌ Файл не найден: {docx_path}")
        sys.exit(1)
    
    analyze_document(docx_path)


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"\n❌ Ошибка: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


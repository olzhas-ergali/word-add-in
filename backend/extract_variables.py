import re
import json
import sys
from pathlib import Path
from uuid import uuid4
from typing import List, Dict, Set

from docx import Document

from app.constants import ENCODING_UTF8


def extract_placeholders_from_docx(docx_path: str) -> Set[str]:
    doc = Document(docx_path)
    placeholders = set()
    
    patterns = [
        r'\{([A-Za-zА-Яа-яЁё0-9_\s]+)\}',           # {VAR}
        r'\{\{([A-Za-zА-Яа-яЁё0-9_\s]+)\}\}',       # {{VAR}}
        r'\[([A-Za-zА-Яа-яЁё0-9_\s]+)\]',           # [VAR]
        r'\[\[([A-Za-zА-Яа-яЁё0-9_\s]+)\]\]',       # [[VAR]]
        r'\$\{([A-Za-zА-Яа-яЁё0-9_\s]+)\}',         # ${VAR}
    ]
    
    print(f"📄 Читаю документ: {docx_path}")
    print(f"   Параграфов: {len(doc.paragraphs)}")
    print(f"   Таблиц: {len(doc.tables)}")
    print()
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        for pattern in patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                placeholders.add(match.strip())
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                for pattern in patterns:
                    matches = re.findall(pattern, text)
                    for match in matches:
                        placeholders.add(match.strip())
    
    try:
        core_props = doc.core_properties
    except:
        pass
    
    return placeholders


def create_variable_database(placeholders: Set[str], document_name: str) -> List[Dict]:
    variables = []
    
    for placeholder in sorted(placeholders):
        variable = {
            "id": str(uuid4()),
            "name": placeholder,
            "display_name": placeholder.replace("_", " ").title(),
            "document": document_name,
            "table": None,  # Заполнить вручную
            "field": None,  # Заполнить вручную
            "value": "",    # Будет заполнено из БД
            "description": "",
            "required": True,
            "data_type": "string"
        }
        variables.append(variable)
    
    return variables


def generate_mapping_template(variables: List[Dict]) -> str:
    template = "# Маппинг переменных на поля базы данных\n"
    template += "# Заполните поля 'table' и 'field' для каждой переменной\n\n"
    
    for var in variables:
        template += f"# {var['name']}\n"
        template += f"# UUID: {var['id']}\n"
        template += f"# Таблица: _______\n"
        template += f"# Поле: _______\n"
        template += f"# Тип данных: string / number / date / boolean\n"
        template += "\n"
    
    return template


def save_to_json(variables: List[Dict], output_path: str):
    with open(output_path, 'w', encoding=ENCODING_UTF8) as f:
        json.dump(variables, f, ensure_ascii=False, indent=2)
    print(f"✅ Сохранено в: {output_path}")


def save_to_sql(variables: List[Dict], output_path: str, document_id: str):
    sql = f"""-- SQL скрипт для создания переменных документа
-- Документ: {document_id}
-- Дата: {Path().absolute()}

-- Таблица документов
INSERT INTO documents (id, name, file_name, created_at) 
VALUES ('{document_id}', 'ДДУ Шымкент', 'ДДУ Шымкент.docx', NOW());

-- Таблица переменных
"""
    
    for var in variables:
        sql += f"""
INSERT INTO document_variables (id, document_id, name, display_name, table_name, field_name, required, data_type)
VALUES (
    '{var['id']}', 
    '{document_id}',
    '{var['name']}',
    '{var['display_name']}',
    NULL,  -- TODO: заполнить table_name
    NULL,  -- TODO: заполнить field_name
    {str(var['required']).lower()},
    '{var['data_type']}'
);
"""
    
    with open(output_path, 'w', encoding=ENCODING_UTF8) as f:
        f.write(sql)
    print(f"✅ SQL скрипт сохранен: {output_path}")


def generate_python_dict(variables: List[Dict]) -> str:
    code = "# Python словарь переменных\n"
    code += "DOCUMENT_VARIABLES = {\n"
    
    for var in variables:
        code += f"    '{var['name']}': {{\n"
        code += f"        'id': '{var['id']}',\n"
        code += f"        'display_name': '{var['display_name']}',\n"
        code += f"        'table': None,  # TODO\n"
        code += f"        'field': None,  # TODO\n"
        code += f"    }},\n"
    
    code += "}\n"
    return code


def main():
    docx_path = "../ДДУ Шымкент.docx"
    
    if not Path(docx_path).exists():
        print(f"❌ Файл не найден: {docx_path}")
        print(f"   Убедитесь, что файл находится в: {Path(docx_path).absolute()}")
        sys.exit(1)
    
    print("=" * 60)
    print("🔍 ИЗВЛЕЧЕНИЕ ПЕРЕМЕННЫХ ИЗ WORD ДОКУМЕНТА")
    print("=" * 60)
    print()
    
    placeholders = extract_placeholders_from_docx(docx_path)
    
    print(f"✅ Найдено переменных: {len(placeholders)}")
    print()
    
    if not placeholders:
        print("⚠️  Переменные не найдены в документе!")
        print("   Проверьте, что документ содержит переменные в формате:")
        print("   {VARIABLE}, {{VARIABLE}}, [VARIABLE], или ${VARIABLE}")
        return
    
    print("📋 Найденные переменные:")
    print("-" * 60)
    for i, var in enumerate(sorted(placeholders), 1):
        print(f"{i:3d}. {var}")
    print()
    
    document_id = str(uuid4())
    document_name = "ДДУ Шымкент"
    variables = create_variable_database(placeholders, document_name)
    
    output_dir = Path("variables_export")
    output_dir.mkdir(exist_ok=True)
    
    print("💾 Сохранение результатов...")
    print()
    
    json_path = output_dir / "variables.json"
    save_to_json(variables, str(json_path))
    
    sql_path = output_dir / "variables.sql"
    save_to_sql(variables, str(sql_path), document_id)
    
    mapping_path = output_dir / "mapping_template.txt"
    mapping_template = generate_mapping_template(variables)
    with open(mapping_path, 'w', encoding=ENCODING_UTF8) as f:
        f.write(mapping_template)
    print(f"✅ Шаблон маппинга: {mapping_path}")
    
    py_path = output_dir / "variables_dict.py"
    py_code = generate_python_dict(variables)
    with open(py_path, 'w', encoding=ENCODING_UTF8) as f:
        f.write(py_code)
    print(f"✅ Python код: {py_path}")
    
    print()
    print("=" * 60)
    print("✅ ГОТОВО!")
    print("=" * 60)
    print()
    print(f"📁 Все файлы сохранены в: {output_dir.absolute()}")
    print()
    print("📝 Следующие шаги:")
    print("   1. Откройте variables.json и проверьте переменные")
    print("   2. Заполните mapping_template.txt (укажите таблицы и поля БД)")
    print("   3. Выполните variables.sql в вашей базе данных")
    print("   4. Обновите ваш API для использования этих переменных")
    print()


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"\n❌ Ошибка: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

